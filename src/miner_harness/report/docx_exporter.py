"""DocxReportExporter — relatório técnico DOCX do ProspectionReport.

Gera documento Word estruturado compatível com due diligence e
relatórios JORC-preliminares, a partir de um ProspectionReport.

Seções:
  1. Sumário Executivo
  2. Alvos Identificados (tabela)
  3. Justificativas por Alvo
  4. Análise por Etapa
  5. Lacunas de Dados Consolidadas
  6. Limitações e Ressalvas
  7. Referências de Dados

Ferramenta: python-docx (puro Python, sem dependências de sistema).

Ref: PRD-002 F9
"""

from __future__ import annotations

from datetime import datetime
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from pathlib import Path

    from miner_harness.core.types import MineralTarget, ProspectionReport, StepResult

# Mapeamento human-readable para cada AnalysisStep
_STEP_LABELS: dict[str, str] = {
    "tectonic_history": "1. História Tectônica",
    "structural_architecture": "2. Arquitetura Estrutural",
    "magmatic_fertility": "3. Fertilidade Magmática",
    "indirect_evidence": "4. Evidências Indiretas",
    "total_integration": "5. Integração Total",
}

_CONFIDENCE_PT: dict[str, str] = {
    "high": "Alta",
    "medium": "Média",
    "low": "Baixa",
    "insufficient": "Insuficiente",
}

_DATA_SOURCES_PT: list[str] = [
    "GeoSGB/CPRM — Ocorrências Minerais, Gravimetria, Geoquímica, Geocronologia, "
    "Litoestratigrafia, Aerogeofísica, Furos de Sondagem",
    "ANM/SIGMINE — Cadastro de Concessões Minerárias",
    "USGS Earthquake Hazards Program — Catálogo de Eventos Sísmicos",
]


class DocxReportExporter:
    """Exporta ProspectionReport para arquivo DOCX estruturado.

    Gera documento Word com sumário executivo, tabela de alvos,
    justificativas por alvo, análise por etapa, lacunas e ressalvas.

    Usage:
        exporter = DocxReportExporter()
        path = exporter.export(report, Path("relatorio.docx"))
    """

    def export(self, report: ProspectionReport, path: Path) -> Path:
        """Gera arquivo DOCX e retorna o path gravado.

        Args:
            report: Relatório de prospecção validado.
            path: Caminho de saída (será criado com parents=True).

        Returns:
            Path do arquivo gerado.

        Raises:
            ImportError: Se python-docx não estiver instalado.
        """
        try:
            from docx import Document  # noqa: PLC0415
        except ImportError as exc:  # pragma: no cover
            msg = (
                "python-docx é necessário para exportar DOCX. Instale com: pip install python-docx"
            )
            raise ImportError(msg) from exc

        doc = Document()

        self._add_title_block(doc, report)
        self._add_executive_summary(doc, report)
        self._add_targets_table(doc, report)
        self._add_target_justifications(doc, report)
        self._add_step_analysis(doc, report)
        self._add_data_gaps(doc, report)
        self._add_caveats(doc, report)
        self._add_data_references(doc, report)

        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        return path

    # ------------------------------------------------------------------
    # Seções privadas
    # ------------------------------------------------------------------

    def _add_title_block(self, doc: Any, report: ProspectionReport) -> None:
        """Página de título com região, data e metadados."""
        from docx.shared import Pt, RGBColor  # noqa: PLC0415

        title = doc.add_heading("RELATÓRIO TÉCNICO DE PROSPECÇÃO MINERAL", level=0)
        title.alignment = 1  # CENTER

        sub = doc.add_paragraph(f"Região: {report.region_name}")
        sub.alignment = 1
        sub.runs[0].font.size = Pt(14)
        sub.runs[0].font.bold = True

        date_str = (
            report.analysis_date.strftime("%d/%m/%Y %H:%M UTC")
            if isinstance(report.analysis_date, datetime)
            else str(report.analysis_date)
        )
        meta = doc.add_paragraph()
        meta.alignment = 1
        meta.add_run(f"Data: {date_str}  |  Modelo: {report.model_used}")
        meta.runs[0].font.size = Pt(10)
        meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        bbox = report.bbox
        bbox_str = (
            f"BBox: {bbox.lon_min:.4f}, {bbox.lat_min:.4f} → {bbox.lon_max:.4f}, {bbox.lat_max:.4f}"
        )
        meta2 = doc.add_paragraph()
        meta2.alignment = 1
        meta2.add_run(
            f"Qualidade dos dados: {report.data_quality_score:.0%}  |  "
            f"Duração: {report.total_duration_ms / 1000:.1f}s  |  {bbox_str}"
        )
        meta2.runs[0].font.size = Pt(9)
        meta2.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()  # espaço

    def _add_executive_summary(self, doc: Any, report: ProspectionReport) -> None:
        """Seção 1 — Sumário Executivo."""
        doc.add_heading("1. Sumário Executivo", level=1)
        doc.add_paragraph(report.integrated_summary or "Não disponível.")

    def _add_targets_table(self, doc: Any, report: ProspectionReport) -> None:
        """Seção 2 — Tabela de Alvos Identificados."""
        doc.add_heading("2. Alvos Identificados", level=1)

        if not report.targets:
            doc.add_paragraph("Nenhum alvo identificado nesta análise.")
            return

        n_cols = 7
        table = doc.add_table(rows=1, cols=n_cols)
        table.style = "Table Grid"

        headers = [
            "Prioridade",
            "Nome",
            "Sistema Mineral",
            "Commodities",
            "Confiança",
            "Longitude / Latitude",
            "Raio (km)",
        ]
        hdr_row = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_row[i].text = h
            # negrito no cabeçalho
            for para in hdr_row[i].paragraphs:
                for run in para.runs:
                    run.bold = True

        for t in report.targets:
            row = table.add_row().cells
            row[0].text = f"P{t.priority}"
            row[1].text = t.name
            row[2].text = t.mineral_system
            row[3].text = ", ".join(t.commodities)
            row[4].text = _CONFIDENCE_PT.get(t.confidence.value, t.confidence.value)
            row[5].text = f"{t.longitude:.4f}, {t.latitude:.4f}"
            row[6].text = f"{t.radius_km:.1f}"

        doc.add_paragraph()

    def _add_target_justifications(self, doc: Any, report: ProspectionReport) -> None:
        """Seção 3 — Justificativas por Alvo."""
        doc.add_heading("3. Justificativas por Alvo", level=1)

        if not report.targets:
            doc.add_paragraph("Nenhum alvo a detalhar.")
            return

        for i, t in enumerate(report.targets, 1):
            self._add_target_section(doc, i, t)

    def _add_target_section(self, doc: Any, idx: int, t: MineralTarget) -> None:
        doc.add_heading(f"3.{idx}. {t.name}", level=2)

        meta = doc.add_paragraph()
        meta.add_run(
            f"Sistema Mineral: {t.mineral_system}  |  "
            f"Commodities: {', '.join(t.commodities)}  |  "
            f"Confiança: {_CONFIDENCE_PT.get(t.confidence.value, t.confidence.value)}  |  "
            f"Prioridade: P{t.priority}"
        )
        meta.runs[0].font.bold = False

        coord_p = doc.add_paragraph()
        coord_p.add_run(
            f"Coordenadas: {t.longitude:.5f}°, {t.latitude:.5f}°  |  "
            f"Raio de investigação: {t.radius_km:.1f} km"
        )

        p_rat = doc.add_paragraph()
        p_rat.add_run("Justificativa:  ").bold = True
        p_rat.add_run(t.rationale)

        if t.recommended_followup:
            doc.add_paragraph().add_run("Follow-up recomendado:").bold = True
            for item in t.recommended_followup:
                doc.add_paragraph(item, style="List Bullet")

        doc.add_paragraph()

    def _add_step_analysis(self, doc: Any, report: ProspectionReport) -> None:
        """Seção 4 — Análise por Etapa."""
        doc.add_heading("4. Análise por Etapa", level=1)

        if not report.steps:
            doc.add_paragraph("Nenhuma etapa registrada.")
            return

        for i, step in enumerate(report.steps, 1):
            self._add_step_section(doc, i, step)

    def _add_step_section(self, doc: Any, idx: int, step: StepResult) -> None:
        label = _STEP_LABELS.get(step.step.value, step.step.value)
        doc.add_heading(f"4.{idx}. {label}", level=2)

        meta = doc.add_paragraph()
        conf_pt = _CONFIDENCE_PT.get(step.confidence.value, step.confidence.value)
        meta.add_run(
            f"Agente: {step.agent}  |  Confiança: {conf_pt}  |  Duração: {step.duration_ms} ms"
        )

        p_sum = doc.add_paragraph()
        p_sum.add_run("Síntese:  ").bold = True
        p_sum.add_run(step.summary)

        if step.findings:
            doc.add_paragraph().add_run("Achados:").bold = True
            for f in step.findings:
                doc.add_paragraph(f, style="List Bullet")

        if step.data_sources_used:
            doc.add_paragraph().add_run("Fontes de dados:").bold = True
            for s in step.data_sources_used:
                doc.add_paragraph(s, style="List Bullet")

        if step.data_gaps:
            doc.add_paragraph().add_run("Lacunas de dados:").bold = True
            for g in step.data_gaps:
                doc.add_paragraph(g, style="List Bullet")

        if step.calibration_note:
            p = doc.add_paragraph()
            p.add_run("Nota de calibração: ").bold = True
            p.add_run(step.calibration_note).italic = True

        doc.add_paragraph()

    def _add_data_gaps(self, doc: Any, report: ProspectionReport) -> None:
        """Seção 5 — Lacunas de Dados Consolidadas."""
        doc.add_heading("5. Lacunas de Dados Consolidadas", level=1)

        # Deduplica gaps de todos os steps
        seen: set[str] = set()
        all_gaps: list[str] = []
        for step in report.steps:
            for gap in step.data_gaps:
                key = gap.lower().strip()
                if key not in seen:
                    seen.add(key)
                    all_gaps.append(gap)

        if not all_gaps:
            doc.add_paragraph("Nenhuma lacuna de dados identificada.")
            return

        for gap in all_gaps:
            doc.add_paragraph(gap, style="List Bullet")

        doc.add_paragraph()

    def _add_caveats(self, doc: Any, report: ProspectionReport) -> None:
        """Seção 6 — Limitações e Ressalvas."""
        doc.add_heading("6. Limitações e Ressalvas", level=1)

        if report.missing_sources or report.bbox_filtered_sources:
            doc.add_paragraph().add_run("Fontes indisponíveis ou sem dados na área:").bold = True
            for s in report.missing_sources:
                doc.add_paragraph(f"{s} — sem dados retornados", style="List Bullet")
            for s in report.bbox_filtered_sources:
                doc.add_paragraph(
                    f"{s} — dados disponíveis mas fora do bbox de análise", style="List Bullet"
                )
            doc.add_paragraph()

        if report.diversity_removed_count > 0:
            doc.add_paragraph(
                f"Nota de diversidade: {report.diversity_removed_count} alvo(s) suprimido(s) "
                "por estarem a menos de 15 km de alvos de maior prioridade "
                "(critério de diversidade espacial mínima).",
                style="List Bullet",
            )

        if report.caveats:
            doc.add_paragraph().add_run("Ressalvas adicionais:").bold = True
            for c in report.caveats:
                doc.add_paragraph(c, style="List Bullet")
        else:
            doc.add_paragraph("Nenhuma ressalva adicional registrada.")

        # Aviso JORC
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        disclaimer.add_run(
            "AVISO: Este relatório é de natureza exploratória e não constitui "
            "estimativa de recursos minerais conforme JORC Code ou NI 43-101. "
            "As interpretações são baseadas em dados públicos e modelos de linguagem, "
            "requerendo confirmação por trabalhos de campo e sondagem."
        )
        disclaimer.runs[0].font.italic = True

    def _add_data_references(self, doc: Any, report: ProspectionReport) -> None:
        """Seção 7 — Referências de Dados."""
        doc.add_heading("7. Referências de Dados", level=1)

        for ref in _DATA_SOURCES_PT:
            doc.add_paragraph(ref, style="List Bullet")

        doc.add_paragraph()
        model_ref = doc.add_paragraph()
        model_ref.add_run("Modelo de linguagem: ").bold = True
        model_ref.add_run(report.model_used)

        doc.add_paragraph()
        gen_p = doc.add_paragraph()
        gen_p.add_run(
            "Relatório gerado automaticamente pelo miner-harness "
            "(https://github.com/VictorAMM/miner-harness)"
        )
        gen_p.runs[0].font.italic = True
