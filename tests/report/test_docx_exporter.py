"""Testes do DocxReportExporter — geração de relatório técnico DOCX.

Ref: PRD-002 F9
"""

from __future__ import annotations

from datetime import datetime, timezone  # noqa: UP017
from typing import TYPE_CHECKING

import pytest

from miner_harness.core.types import (
    AnalysisStep,
    BoundingBox,
    Confidence,
    MineralTarget,
    ProspectionReport,
    StepResult,
)
from miner_harness.report import DocxReportExporter

if TYPE_CHECKING:
    from pathlib import Path

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

docx = pytest.importorskip("docx", reason="python-docx não instalado")


@pytest.fixture()
def sample_target() -> MineralTarget:
    return MineralTarget(
        name="Alvo Serra do Rabo",
        longitude=-49.85,
        latitude=-6.05,
        radius_km=4.5,
        commodities=["Au", "Cu"],
        mineral_system="IOCG",
        confidence=Confidence.HIGH,
        priority=1,
        rationale="Anomalia gravimétrica marcante associada a alteração hidrotermal.",
        recommended_followup=["Sondagem diamantada", "Levantamento IP"],
    )


@pytest.fixture()
def sample_target_low() -> MineralTarget:
    return MineralTarget(
        name="Alvo Sudeste",
        longitude=-50.10,
        latitude=-6.80,
        radius_km=8.0,
        commodities=["Ni"],
        mineral_system="Intrusão Máfica",
        confidence=Confidence.LOW,
        priority=3,
        rationale="Anomalia geoquímica de Ni sem confirmação estrutural.",
        recommended_followup=["Amostragem de solo"],
    )


@pytest.fixture()
def sample_step() -> StepResult:
    return StepResult(
        step=AnalysisStep.TECTONIC_HISTORY,
        agent="structural_geologist",
        summary="Região integra o Cráton Amazônico, com expressivo magmatismo Sideria no.",
        findings=["Evento transamazônico 2.1 Ga", "Domínio Carajás bem delimitado"],
        confidence=Confidence.HIGH,
        data_sources_used=["GeoSGB/litoestratigrafia", "GeoSGB/geocronologia"],
        data_gaps=["Dados de geocronologia ausentes no setor leste"],
        raw_reasoning="Raciocínio detalhado do agente.",
        duration_ms=3800,
    )


@pytest.fixture()
def sample_step2() -> StepResult:
    return StepResult(
        step=AnalysisStep.INDIRECT_EVIDENCE,
        agent="geochemist",
        summary="Anomalias de Cu-Au na geoquímica de sedimento de corrente.",
        findings=["CF(Cu) > 3 em 5 amostras", "Correlação Cu-Au positiva"],
        confidence=Confidence.MEDIUM,
        data_sources_used=["GeoSGB/geoquimica"],
        data_gaps=["Dados de geocronologia ausentes no setor leste"],  # duplicado — testa dedup
        raw_reasoning="Análise geoquímica detalhada.",
        duration_ms=5100,
    )


@pytest.fixture()
def sample_report(
    sample_target: MineralTarget,
    sample_target_low: MineralTarget,
    sample_step: StepResult,
    sample_step2: StepResult,
) -> ProspectionReport:
    return ProspectionReport(
        region_name="Carajás Sul",
        bbox=BoundingBox(lon_min=-51.5, lat_min=-7.0, lon_max=-49.0, lat_max=-5.0),
        analysis_date=datetime(2026, 5, 22, 9, 0, 0, tzinfo=timezone.utc),
        steps=[sample_step, sample_step2],
        targets=[sample_target, sample_target_low],
        integrated_summary="Região com alto potencial para sistemas IOCG. Prioridade P1.",
        caveats=["Dados gravimétricos com cobertura parcial"],
        data_quality_score=0.82,
        total_duration_ms=38000,
        model_used="qwen3:8b",
        missing_sources=["GeoSGB/aerogeofisica"],
        bbox_filtered_sources=["GeoSGB/geocronologia"],
    )


@pytest.fixture()
def minimal_report() -> ProspectionReport:
    """Relatório mínimo sem targets, steps ou extras."""
    return ProspectionReport(
        region_name="Região Teste",
        bbox=BoundingBox(lon_min=-52.0, lat_min=-8.0, lon_max=-50.0, lat_max=-6.0),
        analysis_date=datetime(2026, 5, 1, 0, 0, 0, tzinfo=timezone.utc),
        steps=[],
        targets=[],
        integrated_summary="",
        caveats=[],
        data_quality_score=0.0,
        total_duration_ms=0,
        model_used="qwen3:8b",
    )


# ---------------------------------------------------------------------------
# TestDocxExporterExport — geração de arquivo
# ---------------------------------------------------------------------------


class TestDocxExporterExport:
    def test_export_creates_file(self, sample_report: ProspectionReport, tmp_path: Path) -> None:
        path = tmp_path / "relatorio.docx"
        exporter = DocxReportExporter()
        result = exporter.export(sample_report, path)
        assert result == path
        assert path.exists()
        assert path.stat().st_size > 0

    def test_export_creates_parent_dirs(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "subdir" / "deep" / "report.docx"
        DocxReportExporter().export(sample_report, path)
        assert path.exists()

    def test_exported_file_is_valid_docx(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        """Verifica que o arquivo é um DOCX válido (abrível pelo python-docx)."""
        from docx import Document

        path = tmp_path / "relatorio.docx"
        DocxReportExporter().export(sample_report, path)
        doc = Document(str(path))
        assert len(doc.paragraphs) > 0

    def test_export_minimal_report(self, minimal_report: ProspectionReport, tmp_path: Path) -> None:
        """Relatório sem targets/steps não deve lançar exceção."""
        path = tmp_path / "minimal.docx"
        DocxReportExporter().export(minimal_report, path)
        assert path.exists()


# ---------------------------------------------------------------------------
# TestDocxContent — conteúdo do documento
# ---------------------------------------------------------------------------


class TestDocxContent:
    def _get_full_text(self, path: Path) -> str:
        """Extrai todo o texto de um DOCX como string única."""
        from docx import Document

        doc = Document(str(path))
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            paragraphs.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        paragraphs.append(para.text)
        return "\n".join(paragraphs)

    def test_region_name_in_document(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "Carajás Sul" in text

    def test_title_in_document(self, sample_report: ProspectionReport, tmp_path: Path) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "RELATÓRIO TÉCNICO" in text

    def test_executive_summary_section(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "Sumário Executivo" in text
        assert "alto potencial" in text

    def test_targets_table_contains_target_names(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "Alvo Serra do Rabo" in text
        assert "Alvo Sudeste" in text

    def test_targets_table_contains_commodities(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "Au" in text
        assert "Ni" in text

    def test_targets_table_contains_mineral_system(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "IOCG" in text

    def test_target_justifications_section(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "Justificativas" in text
        assert "anomalia gravimétrica" in text.lower()

    def test_target_followup_in_justifications(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "Sondagem diamantada" in text

    def test_step_analysis_section(self, sample_report: ProspectionReport, tmp_path: Path) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "Análise por Etapa" in text
        assert "História Tectônica" in text

    def test_step_findings_in_document(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "Evento transamazônico" in text

    def test_data_gaps_section(self, sample_report: ProspectionReport, tmp_path: Path) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "Lacunas de Dados" in text
        assert "geocronologia" in text.lower()

    def test_data_gaps_deduplicated(self, sample_report: ProspectionReport, tmp_path: Path) -> None:
        """Gap idêntico em dois steps deve aparecer uma única vez na seção consolidada."""
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        # Conta quantas vezes a lacuna aparece na seção consolidada
        text = self._get_full_text(path)
        gap = "Dados de geocronologia ausentes no setor leste"
        # O gap pode aparecer nas sub-seções dos steps (2x) + 1x na consolidada
        # Mas na seção consolidada deve aparecer apenas 1 vez
        parts = text.split("5. Lacunas de Dados Consolidadas")
        if len(parts) > 1:
            consolidated_section = parts[1]
            assert consolidated_section.count(gap) == 1

    def test_caveats_section(self, sample_report: ProspectionReport, tmp_path: Path) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "Limitações" in text
        assert "cobertura parcial" in text

    def test_missing_sources_in_caveats(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "aerogeofisica" in text

    def test_bbox_filtered_sources_in_caveats(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "geocronologia" in text

    def test_jorc_disclaimer_present(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "JORC" in text

    def test_data_references_section(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "Referências" in text
        assert "GeoSGB" in text
        assert "ANM" in text
        assert "USGS" in text

    def test_model_name_in_references(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "qwen3:8b" in text

    def test_miner_harness_footer_in_references(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "miner-harness" in text


# ---------------------------------------------------------------------------
# TestDocxEdgeCases — casos extremos
# ---------------------------------------------------------------------------


class TestDocxEdgeCases:
    def _get_full_text(self, path: Path) -> str:
        from docx import Document

        doc = Document(str(path))
        parts: list[str] = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(p.text for p in cell.paragraphs)
        return "\n".join(parts)

    def test_no_targets_message(self, minimal_report: ProspectionReport, tmp_path: Path) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(minimal_report, path)
        text = self._get_full_text(path)
        assert "Nenhum alvo" in text

    def test_no_steps_message(self, minimal_report: ProspectionReport, tmp_path: Path) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(minimal_report, path)
        text = self._get_full_text(path)
        assert "Nenhuma etapa" in text

    def test_confidence_labels_pt(self, sample_report: ProspectionReport, tmp_path: Path) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        # Alta e Baixa devem aparecer (targets têm HIGH e LOW)
        assert "Alta" in text
        assert "Baixa" in text

    def test_priority_labels_in_table(
        self, sample_report: ProspectionReport, tmp_path: Path
    ) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "P1" in text
        assert "P3" in text

    def test_coordinates_in_table(self, sample_report: ProspectionReport, tmp_path: Path) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        # Coordenadas do target P1: -49.85, -6.05
        assert "-49.85" in text or "-49.8500" in text

    def test_bbox_in_title_block(self, sample_report: ProspectionReport, tmp_path: Path) -> None:
        path = tmp_path / "r.docx"
        DocxReportExporter().export(sample_report, path)
        text = self._get_full_text(path)
        assert "-51.5" in text or "-51.5000" in text
